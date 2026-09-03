"""Validate a .docx file against the 9 thể thức components of NĐ 30/2020/NĐ-CP.

Usage:
    python validate_thethuc.py <file.docx>

Output: a checklist with ✓ / ⚠ / ✗ per item, plus a summary line.

Note: this is a heuristic check. Manual review is still required for things
like dấu, chữ ký, font sizes (python-docx doesn't reliably read font sizes
from style inheritance).

Rules nguồn: tri-thuc-template/rules/the-thuc.yaml (qua scripts/rules_loader.py).
Nếu YAML thiếu/lỗi → fallback hardcoded defaults dưới đây (backward compat).
"""
import re
import sys
from pathlib import Path

from _common import slugify_vn  # noqa: F401
from docx import Document
from rules_loader import load_rules

OK = "✓"
WARN = "⚠"
FAIL = "✗"


# =====================================================================
# Hardcoded fallback — dùng khi YAML thiếu (offline first-run, etc.)
# =====================================================================
_FALLBACK = {
    "placeholder_pattern": r"\?\?\?|<[^>]+>|\[placeholder\]",
    "bieu_mau_noi_bo_pattern":
        r"PHIẾU\s+(GHI\s+Ý\s+KIẾN|BIỂU\s+QUYẾT|THẨM\s+ĐỊNH|LẤY\s+Ý\s+KIẾN)",
    "le_chuan": {"left_cm": 3, "right_cm": 2, "top_cm": 2, "bottom_cm": 2,
                 "ok_msg": "Lề 3-2-2-2 khớp NĐ30",
                 "fail_template": "Lề hiện tại {actual} — NĐ30 quy định trái 3, phải 2, trên 2, dưới 2 cm"},
    "checks": {
        "quoc_hieu": {
            "must_contain": ["CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
                             "Độc lập", "Tự do", "Hạnh phúc"],
            "ok_msg": "Quốc hiệu + Tiêu ngữ",
            "fail_msg": "Quốc hiệu + Tiêu ngữ — thiếu hoặc sai",
        },
        "co_quan": {
            "upper_line_keywords_regex":
                r"(UBND|ỦY BAN|BỘ|SỞ|HĐND|VĂN PHÒNG|TỔNG CỤC|CỤC|CHI CỤC|VIỆN|BAN)",
            "min_line_len": 4,
            "ok_msg": "Tên cơ quan ban hành",
            "warn_msg": "Tên cơ quan — không phát hiện được dòng tên cơ quan in hoa",
        },
        "so_van_ban": {
            "skip_if_bieu_mau_noi_bo": True,
            "pattern": r"Số:\s*([0-9]*\s*/[A-ZĐa-z&\-_]+)",
            "has_so_marker": "Số:",
            "bieu_mau_msg": "Số VB — không cần (biểu mẫu nội bộ)",
            "ok_template": "Số văn bản: {match}",
            "empty_warn_template": "Số văn bản — đang trống ({match}) — VPHC sẽ điền sau",
            "no_format_warn": "Có dòng 'Số:' nhưng không match format chuẩn",
            "fail_msg": "Không tìm thấy 'Số:' — kiểm tra lại",
        },
        "ten_loai": {
            "keywords": ["QUYẾT ĐỊNH", "NGHỊ QUYẾT", "BÁO CÁO", "TỜ TRÌNH",
                         "THÔNG BÁO", "KẾ HOẠCH", "CHỈ THỊ", "BIÊN BẢN",
                         "GIẤY MỜI", "PHIẾU GHI Ý KIẾN", "PHIẾU BIỂU QUYẾT",
                         "HƯỚNG DẪN", "KẾT LUẬN"],
            "cong_van_markers": ["V/v", "V/V"],
            "ok_template": "Tên loại: {match}",
            "cong_van_msg": "Tên loại: Công văn (có trích yếu V/v...)",
            "warn_msg": "Không xác định được loại VB",
        },
        "noi_dung": {
            "use_placeholder_pattern": True,
            "ok_msg": "Nội dung — không còn placeholder",
            "fail_template": "Còn placeholder: {matches}",
        },
        "nguoi_ky": {
            "pattern": r"(KT\.|TL\.|TUQ\.)?\s*(GIÁM ĐỐC|CHỦ TỊCH|PHÓ CHỦ TỊCH|"
                       r"CHÁNH VĂN PHÒNG|PHÓ CHÁNH VĂN PHÒNG|TRƯỞNG PHÒNG|"
                       r"PHÓ GIÁM ĐỐC|BỘ TRƯỞNG|THỨ TRƯỞNG|VỤ TRƯỞNG|"
                       r"CỤC TRƯỞNG|VIỆN TRƯỞNG)",
            "ok_template": "Chức vụ người ký: {match}",
            "warn_msg": "Không phát hiện được chức vụ người ký in hoa",
        },
        "dau": {
            "msg": "Dấu / chữ ký số — KHÔNG kiểm tra tự động được, kiểm tra thủ công",
            "status": "warn",
        },
        "noi_nhan": {
            "skip_if_bieu_mau_noi_bo": True,
            "noi_nhan_markers": ["Nơi nhận:", "Nơi nhận :"],
            "luu_markers": ["Lưu:", "Lưu :"],
            "bieu_mau_msg": "Nơi nhận — không cần (biểu mẫu nội bộ)",
            "ok_msg": "Nơi nhận + Lưu",
            "missing_luu_warn": "Có 'Nơi nhận' nhưng thiếu 'Lưu:'",
            "fail_msg": "Thiếu 'Nơi nhận:'",
        },
        "phu_luc": {
            "kem_pattern": r"kèm theo|đính kèm",
            "phu_luc_pattern": r"PHỤ LỤC\s+[IVX]+",
            "has_kem_no_pl_warn": "VB nhắc 'kèm theo' nhưng không tìm thấy PHỤ LỤC trong file",
            "has_pl_ok": "Có Phụ lục",
            "no_kem_no_pl_ok": "Không có phụ lục (không bắt buộc)",
        },
    },
}


def _rules() -> dict:
    """Load YAML (cache→bundled) hoặc fallback hardcode."""
    data = load_rules("the-thuc")
    return data if data else _FALLBACK


def _check_cfg(name: str) -> dict:
    return _rules().get("checks", {}).get(name) or _FALLBACK["checks"][name]


# =====================================================================
# Helpers
# =====================================================================

def collect_all_text(doc):
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def placeholder_re() -> re.Pattern:
    pat = _rules().get("placeholder_pattern") or _FALLBACK["placeholder_pattern"]
    return re.compile(pat, re.IGNORECASE)


def is_bieu_mau_noi_bo(text: str) -> bool:
    """Detect biểu mẫu nội bộ — Phiếu biểu quyết / Phiếu ghi ý kiến / Phiếu thẩm định.
    Loại này KHÔNG có Số VB, KHÔNG có Nơi nhận, vì là form gửi nội bộ."""
    pat = _rules().get("bieu_mau_noi_bo_pattern") or _FALLBACK["bieu_mau_noi_bo_pattern"]
    return bool(re.search(pat, text))


# =====================================================================
# 9 mục check
# =====================================================================

def check_quoc_hieu(text):
    cfg = _check_cfg("quoc_hieu")
    if all(s in text for s in cfg["must_contain"]):
        return OK, cfg["ok_msg"]
    return FAIL, cfg["fail_msg"]


def check_co_quan(text):
    cfg = _check_cfg("co_quan")
    min_len = int(cfg.get("min_line_len", 4))
    upper_lines = [ln for ln in text.split("\n")
                   if ln.strip() and ln == ln.upper() and len(ln) > min_len]
    if any(re.search(cfg["upper_line_keywords_regex"], ln) for ln in upper_lines):
        return OK, cfg["ok_msg"]
    return WARN, cfg["warn_msg"]


def check_so_van_ban(text):
    cfg = _check_cfg("so_van_ban")
    if cfg.get("skip_if_bieu_mau_noi_bo") and is_bieu_mau_noi_bo(text):
        return OK, cfg["bieu_mau_msg"]
    m = re.search(cfg["pattern"], text)
    if m:
        sotxt = m.group(1).strip()
        if sotxt.startswith("/"):
            return WARN, cfg["empty_warn_template"].format(match=m.group(0).strip())
        return OK, cfg["ok_template"].format(match=m.group(0).strip())
    if cfg.get("has_so_marker") and cfg["has_so_marker"] in text:
        return WARN, cfg["no_format_warn"]
    return FAIL, cfg["fail_msg"]


def check_ten_loai(text):
    cfg = _check_cfg("ten_loai")
    found = [k for k in cfg["keywords"] if k in text]
    if found:
        return OK, cfg["ok_template"].format(match=found[0])
    if any(m in text for m in cfg.get("cong_van_markers", [])):
        return OK, cfg["cong_van_msg"]
    return WARN, cfg["warn_msg"]


def check_noi_dung(text):
    cfg = _check_cfg("noi_dung")
    rx = placeholder_re()
    if rx.search(text):
        matches = rx.findall(text)[:3]
        return FAIL, cfg["fail_template"].format(matches=matches)
    return OK, cfg["ok_msg"]


def check_nguoi_ky(text):
    cfg = _check_cfg("nguoi_ky")
    m = re.search(cfg["pattern"], text)
    if m:
        return OK, cfg["ok_template"].format(match=m.group(0).strip())
    return WARN, cfg["warn_msg"]


def check_dau():
    cfg = _check_cfg("dau")
    status_map = {"ok": OK, "warn": WARN, "fail": FAIL}
    return status_map.get(cfg.get("status", "warn"), WARN), cfg["msg"]


def check_noi_nhan(text):
    cfg = _check_cfg("noi_nhan")
    if cfg.get("skip_if_bieu_mau_noi_bo") and is_bieu_mau_noi_bo(text):
        return OK, cfg["bieu_mau_msg"]
    has_nn = any(m in text for m in cfg["noi_nhan_markers"])
    has_luu = any(m in text for m in cfg["luu_markers"])
    if has_nn:
        if has_luu:
            return OK, cfg["ok_msg"]
        return WARN, cfg["missing_luu_warn"]
    return FAIL, cfg["fail_msg"]


def check_phu_luc(text):
    cfg = _check_cfg("phu_luc")
    has_kem = bool(re.search(cfg["kem_pattern"], text, re.IGNORECASE))
    has_phu_luc = bool(re.search(cfg["phu_luc_pattern"], text))
    if has_kem and not has_phu_luc:
        return WARN, cfg["has_kem_no_pl_warn"]
    if has_phu_luc:
        return OK, cfg["has_pl_ok"]
    return OK, cfg["no_kem_no_pl_ok"]


# =====================================================================
# CLI
# =====================================================================

def main():
    if len(sys.argv) < 2:
        print("Usage: python validate_thethuc.py <file.docx>", file=sys.stderr)
        return 1

    path = Path(sys.argv[1]).resolve()
    if not path.is_file():
        print(f"ERROR: file not found: {path}", file=sys.stderr)
        return 1

    doc = Document(str(path))
    text = collect_all_text(doc)

    checks = [
        ("1. Quốc hiệu + Tiêu ngữ", check_quoc_hieu(text)),
        ("2. Tên cơ quan ban hành", check_co_quan(text)),
        ("3. Số/ký hiệu", check_so_van_ban(text)),
        ("4. Tên loại + Trích yếu", check_ten_loai(text)),
        ("5. Nội dung", check_noi_dung(text)),
        ("6. Người ký (chức vụ + tên)", check_nguoi_ky(text)),
        ("7. Dấu/chữ ký số", check_dau()),
        ("8. Nơi nhận + Lưu", check_noi_nhan(text)),
        ("9. Phụ lục", check_phu_luc(text)),
    ]

    print(f"=== Validate thể thức: {path.name} ===\n")
    ok = warn = fail = 0
    for label, (status, detail) in checks:
        print(f"  {status} {label}: {detail}")
        if status == OK: ok += 1
        elif status == WARN: warn += 1
        else: fail += 1

    print(f"\nTổng: {OK}{ok}  {WARN}{warn}  {FAIL}{fail}  /  9 mục")
    if fail > 0:
        return 2
    if warn > 0:
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
