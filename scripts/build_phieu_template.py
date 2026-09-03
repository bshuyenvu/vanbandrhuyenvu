"""Build Phiếu ghi ý kiến template + example .docx — biểu mẫu nội bộ UBND.

LOẠI ĐẶC BIỆT — KHÁC VBHC thông thường:
  - Chỉ 1 cơ quan ban hành (UỶ BAN NHÂN DÂN TỈNH) — KHÔNG có "chủ quản + cấp dưới"
  - Tiêu ngữ "Độc lập - Tự do - Hạnh phúc" đậm, không có gạch chân
  - KHÔNG có Số VB / Nơi nhận / Lưu (biểu mẫu nội bộ)
  - Khối ký nằm trong 1 table 1×2 (cell trái empty, cell phải chứa hết)
  - Thứ tự khối ký: ngày → "THÀNH VIÊN UBND TỈNH" → chỗ ký → chức vụ thực → tên

Sinh:
  - resources/templates/phieu-ghi-y-kien.docx
  - examples/Phieu-bieu-quyet-NQ-KHCN-DMST-Vu-Dinh-Hung-example.docx

Reference: D:\\SKILL_AI\\SoanThaoVB_\\Phieu-bieu-quyet-NQ-KHCN-DMST-Vu-Dinh-Hung (1).docx
"""
from __future__ import annotations

import sys, io
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.resolve()))
from vbhc_doc_builder import (
    Document, setup_page,
    add_header_section,
    add_centered_title_with_underline,
    add_kinh_gui,
    add_run,
    add_bieu_quyet_table,
    set_paragraph_spacing,
    set_table_column_widths,
    remove_table_borders,
    set_table_cell_margins_zero,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _condense_run_chars(run, twips: int = -2):
    rPr = run._element.get_or_add_rPr()
    spacing = rPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        rPr.append(spacing)
    spacing.set(qn('w:val'), str(twips))


def _add_phieu_header(doc, dong_1: str, dong_2: str):
    """Header phiếu — TÁI SỬ DỤNG `add_header_section` (giữ gạch chân ngắn của helper),
    chỉ post-modify để cả 2 dòng cơ quan đều ĐẬM (đặc thù phiếu UBND tỉnh: chỉ 1 cơ
    quan, không có cấu trúc chủ quản + cấp dưới như VB thông thường).
    """
    table = add_header_section(
        doc,
        co_quan_chu_quan=dong_1,
        co_quan_ban_hanh=dong_2,
        quoc_hieu_size_pt=13,
        tieu_ngu_size_pt=14,
        co_quan_size_pt=13,
        # Cell trái 5cm vừa đủ "TỈNH TUYÊN QUANG" 2 dòng. Cell phải 11cm dãn hết
        # phần còn lại cho quốc hiệu (16 - 5 = 11cm).
        left_col_cm=5.0,
        right_col_cm=11.0,
        # Gạch chân dưới "Hạnh phúc" = 100% width chữ (đủ hết chiều rộng dòng).
        qh_underline_pct=1.0,
    )
    # Phiếu UBND: dòng 1 cũng ĐẬM (mặc định helper là không đậm)
    left = table.rows[0].cells[0]
    if left.paragraphs and left.paragraphs[0].runs:
        for run in left.paragraphs[0].runs:
            run.bold = True
    # Nén nhẹ char spacing quốc hiệu để vừa 1 dòng ở 13pt (giống pattern báo cáo)
    right = table.rows[0].cells[1]
    if right.paragraphs and right.paragraphs[0].runs:
        for run in right.paragraphs[0].runs:
            _condense_run_chars(run, twips=-2)
    return table


def _add_phieu_signature_block(doc, *, dia_danh: str, ngay, thang, nam,
                                vai_tro: str, chuc_vu_thuc: str, ho_ten: str):
    """Khối ký phiếu: table 1×2, cell trái empty, cell phải chứa toàn bộ.

    Thứ tự: date (italic) → vai_tro (đậm) → 5 dòng trống → chức vụ thực (đậm) → tên (đậm).
    """
    # Đồng bộ với header table: cell trái 5cm + cell phải 11cm. Cell phải đủ chứa
    # chức vụ dài như "GIÁM ĐỐC SỞ GIÁO DỤC VÀ ĐÀO TẠO" (~8.4cm) + khoảng giãn 2 bên.
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_column_widths(table, [5.0, 11.0])
    remove_table_borders(table)
    set_table_cell_margins_zero(table)

    right = table.rows[0].cells[1]
    right.text = ""

    # Date string
    now = datetime.now()
    thang_val = thang or now.month
    nam_val = nam or now.year
    if isinstance(thang_val, int):
        thang_str = f"{thang_val:02d}"
    else:
        thang_str = str(thang_val)
    nam_str = str(nam_val)

    if ngay:
        ngay_str = f"{int(ngay):02d}" if isinstance(ngay, int) else str(ngay)
        date_text = f"{dia_danh}, ngày {ngay_str} tháng {thang_str} năm {nam_str}"
    else:
        date_text = f"{dia_danh}, ngày        tháng {thang_str} năm {nam_str}"

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p, date_text, italic=True, size_pt=13)

    # Vai trò (THÀNH VIÊN UBND TỈNH)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p, vai_tro, bold=True, size_pt=13)

    # 5 dòng trống chỗ ký
    for _ in range(5):
        p = right.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')

    # Chức vụ thực (đậm in hoa)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p, chuc_vu_thuc, bold=True, size_pt=13)

    # Họ tên
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p, ho_ten, bold=True, size_pt=13)

    return table


def build_phieu(out_path: Path, fill: dict):
    doc = Document()
    setup_page(doc)

    # 1. Header (cơ quan UBND tỉnh + quốc hiệu)
    _add_phieu_header(doc,
                      dong_1=fill["co_quan_dong_1"],
                      dong_2=fill["co_quan_dong_2"])

    # 2. Title "PHIẾU GHI Ý KIẾN THÀNH VIÊN UBND TỈNH" + gạch chân ngắn
    # Gạch chân = 1/3 độ dài chữ tiêu đề (chứ KHÔNG phải 1/3 width content).
    # Estimate: TimesNewRoman 14pt bold uppercase ≈ 0.20cm/char. text_width ≈ N*0.20.
    # underline_width = text_width / 3 → indent_each_side = (16 - underline)/2
    _ten = fill["ten_phieu"]
    _text_width = len(_ten) * 0.20
    _underline = _text_width / 3
    _indent = (16.0 - _underline) / 2
    add_centered_title_with_underline(
        doc, _ten,
        size_pt=14, underline_indent_cm=_indent,
    )

    # 3. Kính gửi (KHÔNG dấu chấm cuối — theo mẫu user)
    add_kinh_gui(doc, fill["kinh_gui"])

    # 4. Empty paragraph
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')

    # 5. Italic note "(Đánh dấu X vào một trong hai ô tại Biểu)"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=6, line_pt=1.0, line_rule='auto')
    add_run(p, fill["ghi_chu"], italic=True, size_pt=12)

    # 6. Bảng biểu quyết
    add_bieu_quyet_table(doc, fill["items"])

    # 7. Lý do + Ý kiến khác
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=12, after_pt=0, line_pt=1.5, line_rule='auto')
    add_run(p, "- Lý do không thông qua (đề nghị nêu rõ): ", size_pt=13)
    add_run(p, fill["ly_do"], size_pt=13)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.5, line_rule='auto')
    add_run(p, "- Ý kiến khác (nếu có): ", size_pt=13)
    add_run(p, fill["y_kien_khac"], size_pt=13)

    # 8. Khối ký — table 1×2, cell trái empty, cell phải chứa toàn bộ
    _add_phieu_signature_block(
        doc,
        dia_danh=fill["dia_danh"],
        ngay=fill.get("ngay") or "",
        thang=fill.get("thang") or "",
        nam=fill.get("nam") or "",
        vai_tro=fill["vai_tro"],
        chuc_vu_thuc=fill["chuc_vu_thuc_te"],
        ho_ten=fill["ho_ten_nguoi_ky"],
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


TEMPLATE_FILL = {
    "co_quan_dong_1": "[CO_QUAN_DONG_1]",   # vd: "UỶ BAN NHÂN DÂN"
    "co_quan_dong_2": "[CO_QUAN_DONG_2]",   # vd: "TỈNH TUYÊN QUANG"
    "ten_phieu": "PHIẾU GHI Ý KIẾN THÀNH VIÊN UBND TỈNH",
    "kinh_gui": "[NOI_NHAN_PHIEU]",
    "ghi_chu": "(Đánh dấu X vào một trong hai ô tại Biểu)",
    "items": [
        {"stt": 1, "noi_dung": "[NOI_DUNG_BIEU_QUYET]",
         "dong_y": False, "khong_dong_y": False},
    ],
    "ly_do": "[LY_DO_KHONG_THONG_QUA]",
    "y_kien_khac": "[Y_KIEN_KHAC]",
    "dia_danh": "[DIA_DANH]",
    "ngay": "",
    "thang": "",
    "nam": "",
    "vai_tro": "THÀNH VIÊN UBND TỈNH",
    "chuc_vu_thuc_te": "[CHUC_VU_THUC_TE]",       # vd: "GIÁM ĐỐC SỞ GIÁO DỤC VÀ ĐÀO TẠO"
    "ho_ten_nguoi_ky": "[HO_TEN_NGUOI_KY]",
}

EXAMPLE_FILL = {
    "co_quan_dong_1": "UỶ BAN NHÂN DÂN",
    "co_quan_dong_2": "TỈNH TUYÊN QUANG",
    "ten_phieu": "PHIẾU GHI Ý KIẾN THÀNH VIÊN UBND TỈNH",
    "kinh_gui": "Văn phòng Ủy ban nhân dân tỉnh",
    "ghi_chu": "(Đánh dấu X vào một trong hai ô tại Biểu)",
    "items": [
        {"stt": 1,
         "noi_dung": (
             "Dự thảo Nghị quyết của Hội đồng nhân dân tỉnh quy định một "
             "số nội dung và mức chi sử dụng ngân sách nhà nước cho hoạt "
             "động khoa học, công nghệ và đổi mới sáng tạo tỉnh Tuyên Quang"
         ),
         "dong_y": True, "khong_dong_y": False},
    ],
    "ly_do": "Không.",
    "y_kien_khac": "Không.",
    "dia_danh": "Tuyên Quang",
    "ngay": "",
    "thang": "",
    "nam": "",
    "vai_tro": "THÀNH VIÊN UBND TỈNH",
    "chuc_vu_thuc_te": "GIÁM ĐỐC SỞ GIÁO DỤC VÀ ĐÀO TẠO",
    "ho_ten_nguoi_ky": "Vũ Đình Hưng",
}


def main():
    skill_root = Path(__file__).parent.parent.resolve()
    template_path = skill_root / "resources" / "templates" / "phieu-ghi-y-kien.docx"
    example_path = (
        skill_root / "examples" /
        "Phieu-bieu-quyet-NQ-KHCN-DMST-Vu-Dinh-Hung-example.docx"
    )

    build_phieu(template_path, TEMPLATE_FILL)
    print(f"[OK] Template: {template_path}")

    build_phieu(example_path, EXAMPLE_FILL)
    print(f"[OK] Example:  {example_path}")


if __name__ == "__main__":
    main()
