"""Build Công văn template + example .docx — chuẩn NĐ30.

Sinh:
  - resources/templates/cong-van.docx — template rỗng có placeholder
  - examples/Cong-van-gop-y-du-thao-TT-HBS-So-GDDT-Tuyen-Quang-example.docx — example
    (lấy từ mẫu user `D:\\SKILL_AI\\SoanThaoVB_\\cong-viec\\0003-...`, đã chuẩn hoá:
       - Đ encoding U+00D0 → U+0110
       - Sửa typo "kính giửi" → "kính gửi"
       - Kính gửi tách paragraph riêng (không gộp vào table header))

Chạy:
    python scripts/build_cong_van_template.py
"""
from __future__ import annotations

import sys, io
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.resolve()))
from vbhc_doc_builder import (
    Document, setup_page,
    add_header_section,
    add_so_vb_and_date_section,
    add_kinh_gui,
    add_body_paragraph,
    add_section_heading,
    add_signature_noi_nhan,
    set_paragraph_spacing,
)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _condense_run_chars(run, twips: int = -2):
    rPr = run._element.get_or_add_rPr()
    spacing = rPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        rPr.append(spacing)
    spacing.set(qn('w:val'), str(twips))


def _zero_spacing_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_paragraph_spacing(p, before_pt=0, after_pt=0,
                                      line_pt=1.0, line_rule='auto')


def build_cong_van(out_path: Path, fill: dict):
    doc = Document()
    setup_page(doc)

    # 1. Header
    header_table = add_header_section(
        doc,
        co_quan_chu_quan=fill["co_quan_chu_quan"],
        co_quan_ban_hanh=fill["co_quan_ban_hanh"],
        quoc_hieu_size_pt=13,
        tieu_ngu_size_pt=14,
        co_quan_size_pt=13,
        left_col_cm=6.5,
        right_col_cm=9.5,
    )
    right_cell = header_table.rows[0].cells[1]
    if right_cell.paragraphs and right_cell.paragraphs[0].runs:
        for run in right_cell.paragraphs[0].runs:
            _condense_run_chars(run, twips=-2)

    # 2. Số/V/v + ngày (helper hỗ trợ Công văn — render V/v ngay dưới Số)
    so_table = add_so_vb_and_date_section(
        doc,
        so_vb=fill["so_vb"],
        ky_hieu=fill["ky_hieu"],
        trich_yeu=fill["trich_yeu_vv"],
        dia_danh=fill["dia_danh"],
        ngay=fill["ngay"],
        thang=fill["thang"],
        nam=fill["nam"],
        is_cong_van=True,
        left_col_cm=6.5,
        right_col_cm=9.5,
    )
    _zero_spacing_in_table(so_table)

    # 3. Kính gửi (paragraph riêng — tách khỏi header table cho chuẩn ND30)
    add_kinh_gui(doc, fill["kinh_gui"])

    # 4. Body — list of (kind, ...)
    for entry in fill["body"]:
        kind = entry[0]
        if kind == "para":
            add_body_paragraph(doc, entry[1])
        elif kind == "heading":
            add_section_heading(doc, entry[2], level=entry[1])

    # 5. Câu kết — kết thúc bằng "./."
    cau_ket = fill["cau_ket"].rstrip().rstrip(".") + "./."
    add_body_paragraph(doc, cau_ket)

    # 6. Khối ký + nơi nhận
    add_signature_noi_nhan(
        doc,
        noi_nhan_items=fill["noi_nhan_items"],
        chuc_vu=fill["chuc_vu_nguoi_ky"],
        nguoi_ky=fill["ho_ten_nguoi_ky"],
        quyen_han=fill["quyen_han_ky"],
        chuc_vu_thay=fill["chuc_vu_thay"],
        phong_viet_tat=fill["don_vi_luu"],
    )
    _zero_spacing_in_table(doc.tables[-1])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


TEMPLATE_FILL = {
    "co_quan_chu_quan": "[TEN_CQ_CHU_QUAN]",
    "co_quan_ban_hanh": "[TEN_CQ_BAN_HANH]",
    "so_vb": "",
    "ky_hieu": "[VIET_TAT_CQ]-[VIET_TAT_PHONG_SOAN]",
    "trich_yeu_vv": "[TRICH_YEU_VV]",
    "dia_danh": "[DIA_DANH]",
    "ngay": "",
    "thang": "",
    "nam": "",
    "kinh_gui": "[CO_QUAN_NHAN]",
    "body": [
        ("para", "[DAN_NHAP_THUC_HIEN]"),
        ("para", "[NOI_DUNG_DAN_NHAP_2]"),
        ("heading", 1, "1. [TEN_PHAN_1]"),
        ("para", "[NOI_DUNG_PHAN_1]"),
        ("heading", 1, "2. [TEN_PHAN_2]"),
        ("para", "[NOI_DUNG_PHAN_2]"),
    ],
    "cau_ket": "[CAU_KET]",
    "noi_nhan_items": [
        "- Như trên;",
        "- [NOI_NHAN_2];",
    ],
    "don_vi_luu": "[Phòng soạn thảo văn bản này]",
    "chuc_vu_nguoi_ky": "[CHUC_VU_NGUOI_KY]",
    "ho_ten_nguoi_ky": "[HO_TEN_NGUOI_KY]",
    "quyen_han_ky": "",
    "chuc_vu_thay": "",
}


# Example: từ CV-gop-y-du-thao-TT-HBS-So-GDDT-Tuyen-Quang.docx (đã sửa Đ encoding +
# typo "kính giửi" → "kính gửi", giữ KT. PHÓ GIÁM ĐỐC ký thay GIÁM ĐỐC)
EXAMPLE_FILL = {
    "co_quan_chu_quan": "UBND TỈNH TUYÊN QUANG",
    "co_quan_ban_hanh": "SỞ GIÁO DỤC VÀ ĐÀO TẠO",
    "so_vb": "",
    "ky_hieu": "SGDĐT-GDTrH",
    "trich_yeu_vv": (
        "góp ý dự thảo Thông tư quy định về quản lý và sử dụng "
        "Học bạ số trong các cơ sở giáo dục phổ thông và cơ sở "
        "giáo dục thường xuyên"
    ),
    "dia_danh": "Tuyên Quang",
    "ngay": "",
    "thang": "",
    "nam": "",
    "kinh_gui": "Bộ Giáo dục và Đào tạo",
    "body": [
        ("para",
         "Thực hiện Công văn số 2250/BGDĐT-GDPT ngày 29/4/2026 của Bộ "
         "Giáo dục và Đào tạo về việc góp ý dự thảo Thông tư quy định "
         "về quản lý và sử dụng Học bạ số."),
        ("para",
         "Sở Giáo dục và Đào tạo tỉnh Tuyên Quang đã tổ chức lấy ý kiến "
         "rộng rãi của 235 cán bộ quản lý, giáo viên, nhân viên các cơ "
         "sở giáo dục và cơ quan quản lý trên địa bàn. Sau khi nghiên "
         "cứu, Sở Giáo dục và Đào tạo có ý kiến góp ý như sau:"),
        ("heading", 1, "1. Góp ý cụ thể từng điều"),
        ("para",
         "1.1. Tại Điều 4, khoản 3, điểm đ: đề nghị bỏ cụm từ \"có "
         "điều kiện khai thác, sử dụng dữ liệu\". Lý do: quy định này "
         "dẫn đến giá trị pháp lý của Học bạ số phụ thuộc vào điều "
         "kiện hạ tầng, không khả thi với các cơ sở giáo dục vùng khó "
         "khăn."),
        ("para",
         "1.2. Tại Điều 7, khoản 1, điểm a: đề nghị bỏ cụm từ \"Mã số "
         "hồ sơ học tập suốt đời (mã người học)\" và sử dụng số định "
         "danh cá nhân. Lý do: số định danh cá nhân đã đầy đủ chức năng "
         "định danh duy nhất, không cần thêm mã song song."),
        ("para",
         "1.3. Tại Điều 11: đề nghị sửa cụm từ \"yêu cầu chỉnh sửa\" "
         "(thông tin trên Học bạ số) thành \"đề nghị rà soát\". Lý do: "
         "việc chỉnh sửa nội dung Học bạ phải qua quy trình thẩm định, "
         "không thể \"yêu cầu\" trực tiếp."),
        ("para",
         "1.4. Tại Điều 17 và Điều 19 (về bảo mật, an toàn thông tin): "
         "đề nghị quy định theo nguyên tắc dẫn chiếu \"thực hiện theo "
         "quy định của pháp luật về an toàn, an ninh mạng\" thay vì "
         "liệt kê biện pháp cụ thể, để bảo đảm cập nhật theo công nghệ."),
        ("para",
         "1.5. Tại các Điều 19, 20, 21 (về trách nhiệm tổ chức thực "
         "hiện): đề nghị phân định rõ hai nhóm trách nhiệm: (i) Quản lý "
         "chuyên môn (Bộ Giáo dục và Đào tạo, Sở Giáo dục và Đào tạo) "
         "và (ii) Vận hành kỹ thuật (đơn vị cung cấp hạ tầng)."),
        ("heading", 1, "2. Đề xuất, kiến nghị"),
        ("para",
         "Để bảo đảm tính khả thi, đồng bộ và sớm hoàn thành lộ trình "
         "triển khai Học bạ số toàn quốc, Sở Giáo dục và Đào tạo tỉnh "
         "Tuyên Quang trân trọng đề nghị:"),
        ("para",
         "2.1. Sớm ban hành hướng dẫn số hóa học bạ giấy đối với học "
         "sinh chưa có Học bạ số ở các năm học, lớp học, cấp học trước "
         "(đặc biệt là học sinh đã tốt nghiệp THCS, THPT)."),
        ("para",
         "2.2. Phối hợp với Bộ Tài chính ban hành hướng dẫn đầy đủ, "
         "chi tiết về kinh phí thực hiện Học bạ số (nội dung chi và "
         "định mức chi đối với chữ ký số, lưu trữ, vận hành)."),
        ("para",
         "2.3. Bổ sung trong Thông tư cơ chế phối hợp, hỗ trợ về chữ "
         "ký số, an toàn thông tin và bảo mật dữ liệu giữa các bên "
         "liên quan (Bộ Giáo dục và Đào tạo, Bộ Công an, đơn vị cung "
         "cấp hạ tầng, các cơ sở giáo dục)."),
    ],
    # Đã sửa typo "giửi" → "gửi" (NĐ30 + chính tả chuẩn)
    "cau_ket": "Sở Giáo dục và Đào tạo trân trọng kính gửi.",
    "noi_nhan_items": [
        "- Như trên;",
        "- Lãnh đạo Sở GDĐT;",
    ],
    "don_vi_luu": "GDTrH",
    "chuc_vu_nguoi_ky": "PHÓ GIÁM ĐỐC",
    "ho_ten_nguoi_ky": "Đinh Thế Hiệp",
    "quyen_han_ky": "KT.",
    "chuc_vu_thay": "GIÁM ĐỐC",
}


def main():
    skill_root = Path(__file__).parent.parent.resolve()
    template_path = skill_root / "resources" / "templates" / "cong-van.docx"
    example_path = (
        skill_root / "examples" /
        "Cong-van-gop-y-du-thao-TT-HBS-So-GDDT-Tuyen-Quang-example.docx"
    )

    build_cong_van(template_path, TEMPLATE_FILL)
    print(f"[OK] Template: {template_path}")

    build_cong_van(example_path, EXAMPLE_FILL)
    print(f"[OK] Example:  {example_path}")


if __name__ == "__main__":
    main()
