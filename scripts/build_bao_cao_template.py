"""Build báo cáo định kỳ tổng hợp template + example .docx — chuẩn NĐ 30/2020/NĐ-CP.

Sinh 2 file:
  - resources/templates/bao-cao.docx — template rỗng có placeholder [KEY]
  - examples/Bao-cao-Quy-I-Nam-2026-So-GDDT-Tuyen-Quang.docx — example đã fill

Quy ước style (theo yêu cầu):
  - Tiêu ngữ "CỘNG HÒA..." cỡ 13pt (đúng NĐ30), nén char spacing -0.1pt nếu cần để 1 dòng.
  - Trích yếu được tách thành nhiều dòng cân đối.
  - Số/ngày: chỉ điền tháng-năm hiện tại, để trống ngày (VPHC điền khi ban hành).
  - Spacing trước/sau:
      * Thân báo cáo (body): trước/sau 6pt
      * Tất cả phần còn lại (tiêu ngữ, số/ngày, trích yếu, ký tên, nơi nhận): 0pt
  - Tất cả tiêu đề trong nội dung thụt đầu dòng 1.1cm.
  - Tiêu đề cấp 2 không nghiêng (đậm). Cấp 3 nghiêng, không đậm.
  - Báo cáo có phần "Căn cứ" / "Thực hiện ..." sau title block.

Chạy:
    python scripts/build_bao_cao_template.py
"""
from __future__ import annotations

import sys, io
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.resolve()))
from vbhc_doc_builder import (
    Document, setup_page,
    add_header_section,
    add_so_vb_and_date_section,
    add_run,
    add_body_paragraph,
    add_section_heading,
    add_signature_noi_nhan,
    set_paragraph_spacing,
    set_paragraph_indent,
    set_paragraph_borders,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# Helpers local — bù vào những gì vbhc_doc_builder không expose
# ============================================================

def _zero_spacing_in_table(table):
    """Force spacing before/after = 0, line = single cho mọi paragraph trong table."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_paragraph_spacing(p, before_pt=0, after_pt=0,
                                      line_pt=1.0, line_rule='auto')


def _condense_run_chars(run, twips: int = -2):
    """Nén char spacing 1 run để đảm bảo 1 dòng. twips=-2 ≈ -0.1pt mỗi ký tự."""
    rPr = run._element.get_or_add_rPr()
    spacing = rPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        rPr.append(spacing)
    spacing.set(qn('w:val'), str(twips))


def _add_title_block_multiline(doc, *, ten_loai: str, trich_yeu_lines):
    """Title block: TÊN LOẠI in hoa + trích yếu có thể nhiều dòng + gạch chân ngắn.

    Spacing before/after = 0 cho mọi paragraph.
    """
    if isinstance(trich_yeu_lines, str):
        trich_yeu_lines = [trich_yeu_lines]

    # before_pt=14pt ≈ 1 dòng cỡ 14 → cách header một dòng trống
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=14, after_pt=0, line_pt=1.0, line_rule='auto')
    add_run(p, ten_loai.upper(), bold=True, size_pt=14)

    for line in trich_yeu_lines:
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1.0, line_rule='auto')
        add_run(p, line, bold=True, size_pt=14)

    # Gạch chân ngắn (~35% bề rộng khối nội dung 16cm)
    content_width_cm = 16.0
    underline_pct = 0.35
    underline_width = content_width_cm * underline_pct
    indent_each_side = (content_width_cm - underline_width) / 2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_indent(p,
                         left_twips=int(indent_each_side * 567),
                         right_twips=int(indent_each_side * 567))
    set_paragraph_spacing(p, before_pt=0, after_pt=0, line_pt=1, line_rule='exact')
    set_paragraph_borders(p,
                          top={'val': 'single', 'sz': 4, 'color': '000000', 'space': 0})


# ============================================================
# Builder chính
# ============================================================

def build_bao_cao(out_path: Path, fill: dict):
    """Build báo cáo định kỳ tổng hợp .docx với fill dict."""
    doc = Document()
    setup_page(doc)

    # 1. Header (cơ quan + quốc hiệu) — quốc hiệu cỡ 13pt + nén char để 1 dòng
    header_table = add_header_section(
        doc,
        co_quan_chu_quan=fill["co_quan_chu_quan"],
        co_quan_ban_hanh=fill["co_quan_ban_hanh"],
        quoc_hieu_size_pt=13,
        tieu_ngu_size_pt=14,
        co_quan_size_pt=13,
        left_col_cm=6.5,
        right_col_cm=9.5,
    )
    # Tiêu ngữ là paragraph[0] cell phải. Nén char -0.1pt để chắc chắn 1 dòng ở 13pt.
    right_cell = header_table.rows[0].cells[1]
    if right_cell.paragraphs and right_cell.paragraphs[0].runs:
        for run in right_cell.paragraphs[0].runs:
            _condense_run_chars(run, twips=-2)

    # 2. Số / địa danh-ngày — spacing 0
    so_table = add_so_vb_and_date_section(
        doc,
        so_vb=fill["so_vb"],
        ky_hieu=fill["ky_hieu"],
        dia_danh=fill["dia_danh"],
        ngay=fill["ngay"],
        thang=fill["thang"],
        nam=fill["nam"],
        is_cong_van=False,
        left_col_cm=6.5,
        right_col_cm=9.5,
    )
    _zero_spacing_in_table(so_table)

    # 3. Title block (BÁO CÁO + trích yếu nhiều dòng)
    _add_title_block_multiline(
        doc,
        ten_loai="BÁO CÁO",
        trich_yeu_lines=fill["trich_yeu_lines"],
    )

    # 4. Phần CĂN CỨ + dẫn nhập (body paragraphs, indent 1.1cm)
    for can_cu in fill["can_cu_list"]:
        if can_cu:
            add_body_paragraph(doc, can_cu)
    if fill.get("dan_nhap"):
        add_body_paragraph(doc, fill["dan_nhap"])

    # 5. Khung 4 phần — heading thụt 1.1cm (default helper)
    add_section_heading(doc, "I. TÌNH HÌNH VÀ KẾT QUẢ THỰC HIỆN", level=1)
    add_section_heading(doc, "1. Công tác chỉ đạo, điều hành", level=2)
    add_body_paragraph(doc, fill["noi_dung_i_1"])
    add_section_heading(doc, "2. Kết quả thực hiện các nhiệm vụ trọng tâm", level=2)
    add_body_paragraph(doc, fill["noi_dung_i_2"])

    add_section_heading(doc, "II. ĐÁNH GIÁ CHUNG", level=1)
    add_section_heading(doc, "1. Ưu điểm", level=2)
    add_body_paragraph(doc, fill["noi_dung_ii_1"])
    add_section_heading(doc, "2. Hạn chế, tồn tại", level=2)
    add_body_paragraph(doc, fill["noi_dung_ii_2"])
    add_section_heading(doc, "3. Nguyên nhân", level=2)
    add_body_paragraph(doc, fill["noi_dung_ii_3"])

    add_section_heading(
        doc,
        f"III. PHƯƠNG HƯỚNG, NHIỆM VỤ TRỌNG TÂM {fill['ky_tiep_theo'].upper()}",
        level=1,
    )
    add_body_paragraph(doc, fill["noi_dung_iii"])

    add_section_heading(doc, "IV. ĐỀ XUẤT, KIẾN NGHỊ", level=1)
    add_body_paragraph(doc, fill["noi_dung_iv"])

    # Câu kết VB + ký hiệu "./." (NĐ30 — kết thúc nội dung VB)
    cau_ket = fill["cau_ket"].rstrip().rstrip(".") + "./."
    add_body_paragraph(doc, cau_ket)

    # 6. Khối ký tên + nơi nhận — spacing = 0 toàn bộ
    add_signature_noi_nhan(
        doc,
        noi_nhan_items=fill["noi_nhan_items"],
        chuc_vu=fill["chuc_vu_nguoi_ky"],
        nguoi_ky=fill["ho_ten_nguoi_ky"],
        quyen_han=fill["quyen_han_ky"],
        chuc_vu_thay=fill["chuc_vu_thay"],
        phong_viet_tat=fill["don_vi_luu"],
    )
    _zero_spacing_in_table(doc.tables[-1])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ============================================================
# Dữ liệu — TEMPLATE (placeholder rỗng)
# ============================================================

TEMPLATE_FILL = {
    "co_quan_chu_quan": "[TEN_CQ_CHU_QUAN]",
    "co_quan_ban_hanh": "[TEN_CQ_BAN_HANH]",
    "ky_hieu": "BC-[VIET_TAT_CQ]",
    "so_vb": "",                # giữ pattern "Số:        /BC-..." cho VPHC điền
    "trich_yeu_lines": ["[TRICH_YEU_DONG_1]", "[TRICH_YEU_DONG_2]"],
    "dia_danh": "[DIA_DANH]",
    "ngay": "",                 # luôn để trống — VPHC điền khi ban hành
    "thang": "",                # rỗng → helper auto fill tháng hiện tại
    "nam": "",                  # rỗng → helper auto fill năm hiện tại
    "can_cu_list": [
        "[CAN_CU_1];",
        "[CAN_CU_2].",
    ],
    "dan_nhap": "[TEN_CQ_BAN_HANH_DANH_TU] báo cáo [PHAM_VI_BAO_CAO] như sau:",
    "ky_tiep_theo": "[KY_TIEP_THEO]",
    "noi_dung_i_1": "[NOI_DUNG_I_1]",
    "noi_dung_i_2": "[NOI_DUNG_I_2]",
    "noi_dung_ii_1": "[NOI_DUNG_II_1]",
    "noi_dung_ii_2": "[NOI_DUNG_II_2]",
    "noi_dung_ii_3": "[NOI_DUNG_II_3]",
    "noi_dung_iii": "[NOI_DUNG_III]",
    "noi_dung_iv": "[NOI_DUNG_IV]",
    "cau_ket": "[CAU_KET]",
    "noi_nhan_items": [
        "- [NOI_NHAN_1];",
        "- [NOI_NHAN_2];",
        "- Lãnh đạo cơ quan;",
        "- [NOI_NHAN_3];",
    ],
    "don_vi_luu": "[Phòng soạn thảo văn bản này]",
    "chuc_vu_nguoi_ky": "[CHUC_VU_NGUOI_KY]",
    "ho_ten_nguoi_ky": "[HO_TEN_NGUOI_KY]",
    "quyen_han_ky": "",
    "chuc_vu_thay": "",
}


# ============================================================
# Dữ liệu — EXAMPLE (Báo cáo Quý I/2026 — Sở GD&ĐT Tuyên Quang)
# Hôm nay 2026-05-10 → tháng/năm hiện tại = 5/2026 (helper auto fill)
# ============================================================

EXAMPLE_FILL = {
    "co_quan_chu_quan": "UBND TỈNH TUYÊN QUANG",
    "co_quan_ban_hanh": "SỞ GIÁO DỤC VÀ ĐÀO TẠO",
    "ky_hieu": "BC-SGDĐT",
    "so_vb": "",
    # Trích yếu chia 2 dòng cân đối, ngắt ở chữ "và" cho tự nhiên
    "trich_yeu_lines": [
        "Công tác giáo dục và đào tạo Quý I năm 2026",
        "và nhiệm vụ trọng tâm Quý II năm 2026",
    ],
    "dia_danh": "Tuyên Quang",
    "ngay": "",        # để trống ngày
    "thang": "",       # auto = tháng hiện tại
    "nam": "",         # auto = năm hiện tại
    "can_cu_list": [
        "Thực hiện Chương trình công tác năm 2026 của Sở Giáo dục và Đào tạo "
        "tỉnh Tuyên Quang;",
        "Thực hiện chỉ đạo của Bộ Giáo dục và Đào tạo và UBND tỉnh Tuyên Quang "
        "về nhiệm vụ năm học 2025-2026.",
    ],
    "dan_nhap": (
        "Sở Giáo dục và Đào tạo báo cáo kết quả công tác Quý I năm 2026 và "
        "nhiệm vụ trọng tâm Quý II năm 2026 như sau:"
    ),
    "ky_tiep_theo": "Quý II năm 2026",
    "noi_dung_i_1": (
        "Trong Quý I năm 2026, Sở Giáo dục và Đào tạo đã ban hành đầy đủ các "
        "văn bản chỉ đạo triển khai nhiệm vụ học kỳ II năm học 2025-2026; tổ "
        "chức quán triệt sâu rộng tới các đơn vị trực thuộc; phối hợp chặt "
        "chẽ với UBND các huyện, thành phố trong công tác quản lý, điều "
        "hành; thực hiện tốt công tác kiểm tra chuyên đề và đột xuất."
    ),
    "noi_dung_i_2": (
        "Hoàn thành tổ chức kỳ thi chọn học sinh giỏi cấp tỉnh năm học "
        "2025-2026 với 12 môn học, có 2.150 học sinh dự thi. Triển khai đại "
        "trà học bạ điện tử ở 100% trường phổ thông, kết nối liên thông với "
        "hệ thống cơ sở dữ liệu ngành. Tổ chức thành công Hội thi giáo viên "
        "dạy giỏi cấp tỉnh bậc tiểu học. Công tác bồi dưỡng giáo viên thực "
        "hiện Chương trình giáo dục phổ thông 2018 tiếp tục được đẩy mạnh."
    ),
    "noi_dung_ii_1": (
        "Công tác lãnh đạo, chỉ đạo bám sát chương trình công tác và kịp "
        "thời tháo gỡ vướng mắc tại cơ sở. Đội ngũ cán bộ, giáo viên ổn "
        "định, có ý thức học tập, bồi dưỡng nâng cao chuyên môn nghiệp vụ. "
        "Việc ứng dụng công nghệ thông tin trong dạy học và quản lý giáo "
        "dục được đẩy mạnh."
    ),
    "noi_dung_ii_2": (
        "Vẫn còn tình trạng thiếu giáo viên ở một số môn học, đặc biệt là "
        "Tin học, Tiếng Anh tại các xã vùng sâu, vùng xa. Cơ sở vật chất, "
        "trang thiết bị dạy học của một số trường còn thiếu so với yêu cầu "
        "Chương trình giáo dục phổ thông 2018."
    ),
    "noi_dung_ii_3": (
        "Nguyên nhân chủ quan do công tác dự báo nhu cầu giáo viên ở một số "
        "đơn vị chưa sát thực tế. Nguyên nhân khách quan do địa bàn rộng, "
        "nhiều xã miền núi đặc biệt khó khăn; nguồn ngân sách hạn chế."
    ),
    "noi_dung_iii": (
        "Tập trung tổ chức tốt kỳ thi tốt nghiệp THPT năm 2026 và kỳ thi "
        "tuyển sinh vào lớp 10 năm học 2026-2027. Tiếp tục triển khai có "
        "hiệu quả Chương trình giáo dục phổ thông 2018 đối với các lớp cuối "
        "cấp. Khắc phục tình trạng thiếu giáo viên; rà soát, sắp xếp mạng "
        "lưới trường lớp phù hợp với điều kiện thực tế. Đẩy mạnh chuyển đổi "
        "số trong giáo dục."
    ),
    "noi_dung_iv": (
        "Đề nghị Bộ Giáo dục và Đào tạo tiếp tục hướng dẫn cụ thể về định "
        "mức giáo viên cho các môn học mới và cơ chế hợp đồng đối với giáo "
        "viên Tin học, Tiếng Anh ở vùng khó khăn. Đề nghị UBND tỉnh quan "
        "tâm bố trí ngân sách đầu tư cơ sở vật chất, trang thiết bị dạy "
        "học cho các trường còn thiếu, ưu tiên xã miền núi và vùng đặc "
        "biệt khó khăn."
    ),
    "cau_ket": (
        "Trên đây là báo cáo kết quả công tác Quý I năm 2026 và phương "
        "hướng, nhiệm vụ trọng tâm Quý II năm 2026 của Sở Giáo dục và Đào "
        "tạo, trân trọng báo cáo."
    ),
    "noi_nhan_items": [
        "- UBND tỉnh (b/c);",
        "- Bộ Giáo dục và Đào tạo (b/c);",
        "- Lãnh đạo Sở;",
        "- Các phòng, ban thuộc Sở;",
        "- Phòng GD&ĐT các huyện, thành phố;",
    ],
    "don_vi_luu": "VP",
    "chuc_vu_nguoi_ky": "GIÁM ĐỐC",
    "ho_ten_nguoi_ky": "Vũ Đình Hưng",
    "quyen_han_ky": "",
    "chuc_vu_thay": "",
}


def main():
    skill_root = Path(__file__).parent.parent.resolve()
    template_path = skill_root / "resources" / "templates" / "bao-cao.docx"
    example_path = (
        skill_root / "examples" /
        "Bao-cao-Quy-I-Nam-2026-So-GDDT-Tuyen-Quang.docx"
    )

    build_bao_cao(template_path, TEMPLATE_FILL)
    print(f"[OK] Template: {template_path}")

    build_bao_cao(example_path, EXAMPLE_FILL)
    print(f"[OK] Example:  {example_path}")


if __name__ == "__main__":
    main()
