"""MCP server for soan-thao-vbhc skill (v1.0 — local thin-MCP).

Exposes 14 tools cho workflow soạn VBHC theo Nghị định 30/2020.

Chạy local stdio (mỗi máy user):
    python server.py

Storage tiers:
  - SKILL    = D:/SKILL_AI/skills/soan-thao-vbhc/   (read-only code + danh-muc-loai-vb)
  - CACHE    = ~/.vbhc/cache/  (templates + rules + code — sync từ cloud KB Hub)
  - ORG      = $VBHC_ORG_DIR (default: ~/.vbhc/org/)
               YAML chung cơ quan: thông tin co quan, người ký, phân công nhiệm vụ
  - USER     = arg in tool calls (work folder per task)

Knowledge sync: chạy `python mcp/bootstrap.py` lần đầu (hoặc dùng PowerShell
installer — xem INSTALL-LOCAL.md). Sau đó `vbhc_sync_knowledge` để pull update.

Client config (stdio):
    {"mcpServers": {"vbhc": {"command": "python", "args": [".../server.py"]}}}

HTTP mode (v0.9) đã bị gỡ — xem MIGRATION-v1.0.md để chuyển sang local.
"""
from __future__ import annotations

import sys
import io
import os
import re
import shutil
import unicodedata
from pathlib import Path
from typing import Any

# Ensure UTF-8 output
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")

# =====================================================================
# 3-tier storage layout
# =====================================================================
# Tier 1 — SKILL: code + danh mục VB chuẩn (read-only cho user)
SKILL_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(SKILL_DIR / "scripts"))

# Tier 2 — ORG: cấu hình chung cơ quan (cơ quan, người ký, phân công NV).
# Resolved from env var VBHC_ORG_DIR. Default = ~/.vbhc/org/ (OS home).
def _resolve_org_dir() -> Path:
    env = os.environ.get("VBHC_ORG_DIR")
    if env:
        return Path(env).expanduser().resolve()
    return (Path.home() / ".vbhc" / "org").resolve()

ORG_DIR = _resolve_org_dir()

# Tier 3 — USER: work folder, passed per-call as arg (parent_dir / work_folder).
# Optional default via VBHC_USER_DIR env var.
def _default_user_dir() -> Path | None:
    env = os.environ.get("VBHC_USER_DIR")
    return Path(env).expanduser().resolve() if env else None

try:
    from mcp.server.fastmcp import FastMCP
except ImportError:
    print(
        "ERROR: 'mcp' Python package not found. Install with:\n"
        "    pip install mcp\n"
        "or:\n"
        "    uv tool install mcp",
        file=sys.stderr,
    )
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("ERROR: 'python-docx' not found. Install: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

# ---- helpers ----

def slugify_vn(text: str) -> str:
    text = text.replace("đ", "d").replace("Đ", "d")
    nfkd = unicodedata.normalize("NFKD", text)
    ascii_text = "".join(c for c in nfkd if not unicodedata.combining(c))
    ascii_text = ascii_text.lower()
    ascii_text = re.sub(r"[^a-z0-9]+", "-", ascii_text)
    return ascii_text.strip("-")


def next_folder_number(parent: Path) -> str:
    parent.mkdir(parents=True, exist_ok=True)
    nums = []
    for child in parent.iterdir():
        if not child.is_dir():
            continue
        m = re.match(r"^(\d{4})-", child.name)
        if m:
            nums.append(int(m.group(1)))
    return f"{(max(nums) + 1) if nums else 1:04d}"


def replace_paragraph_text(p, new_text: str):
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


def cell_set_text(cell, new_text: str):
    if not cell.paragraphs:
        cell.add_paragraph(new_text)
        return
    replace_paragraph_text(cell.paragraphs[0], new_text)
    for p_extra in list(cell.paragraphs[1:]):
        p_extra._element.getparent().remove(p_extra._element)


def search_and_replace_doc(doc, find: str, repl: str) -> int:
    count = 0
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        if find in full:
            replace_paragraph_text(p, full.replace(find, repl))
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if find in cell.text:
                    cell_set_text(cell, cell.text.replace(find, repl))
                    count += 1
    return count


# ---- VB classifier ----
#
# Rules nguồn: tri-thuc-template/rules/loai-vb.yaml (qua scripts/rules_loader.py).
# Hardcode dưới đây là fallback khi YAML thiếu (offline first-run, etc.).
# ----------------------------------------------------------------------

from rules_loader import load_rules  # noqa: E402
sys.path.insert(0, str(Path(__file__).resolve().parent))
import knowledge_client as kc  # noqa: E402


def _resolve_template_path(spec: str) -> Path:
    """Resolve template spec → Path:
       1. Nếu spec là path (có sep hoặc tồn tại trên disk) → resolve tuyệt đối
       2. Nếu spec là slug (vd "bao-cao") → tìm ~/.vbhc/cache/templates/<slug>.docx
       3. Fallback bundled: SKILL_DIR/resources/templates/<slug>.docx
       4. Nếu cache thiếu + có cloud config → auto-pull blocking trước khi raise.
    """
    raw = spec
    if any(sep in raw for sep in ("/", "\\")) or raw.endswith(".docx"):
        # Cho phép cả "bao-cao.docx" → coi như slug nếu không có separator
        if not any(sep in raw for sep in ("/", "\\")):
            slug = raw[:-5] if raw.endswith(".docx") else raw
        else:
            return Path(raw).expanduser().resolve()
    else:
        slug = raw

    cache_p = kc.CACHE_DIR / "templates" / f"{slug}.docx"
    if cache_p.is_file():
        return cache_p

    # Auto-bootstrap: thử pull từ cloud nếu đã config
    if kc.is_configured():
        try:
            kc.ensure_asset("templates", f"{slug}.docx")
            if cache_p.is_file():
                return cache_p
        except kc.KBError:
            pass

    bundled = SKILL_DIR / "resources" / "templates" / f"{slug}.docx"
    if bundled.is_file():
        return bundled

    # Nếu vẫn không có — trả Path(raw) để caller báo lỗi rõ ràng (file not found)
    return Path(raw).expanduser().resolve()

_CLASSIFY_FALLBACK = [
    (r"xin\s+ý\s+kiến|lấy\s+ý\s+kiến|tham\s+gia\s+ý\s+kiến", "Công văn xin ý kiến"),
    (r"phiếu\s+(biểu\s+quyết|ghi\s+ý\s+kiến)", "Phiếu biểu quyết / Phiếu ghi ý kiến"),
    (r"tờ\s+trình|kính\s+trình", "Tờ trình"),
    (r"quyết\s+định.*(thành\s+lập|bổ\s+nhiệm|ban\s+hành)", "Quyết định cá biệt"),
    (r"nghị\s+quyết.*hđnd", "Nghị quyết HĐND (QPPL)"),
    (r"báo\s+cáo", "Báo cáo"),
    (r"thông\s+báo|kết\s+luận\s+của", "Thông báo / Kết luận"),
    (r"kế\s+hoạch.*thực\s+hiện", "Kế hoạch"),
    (r"hướng\s+dẫn", "Hướng dẫn"),
    (r"giấy\s+mời", "Giấy mời"),
    (r"biên\s+bản", "Biên bản"),
    (r"chỉ\s+thị", "Chỉ thị"),
    (r"công\s+văn|công\s+điện", "Công văn"),
]

_AMBIGUOUS_FALLBACK = [
    {
        "pattern": r"góp\s+ý|tham\s+gia\s+ý\s+kiến|phản\s+hồi|trả\s+lời",
        "context_hint": "VB phản hồi/góp ý",
        "forms": [
            {"type": "Báo cáo", "when": "VB nguồn yêu cầu BÁO CÁO, hoặc nội dung dài có cấu trúc đề mục"},
            {"type": "Công văn", "when": "VB nguồn yêu cầu CÔNG VĂN, hoặc nội dung ngắn (1-2 trang)"},
        ],
    },
    {
        "pattern": r"phúc\s+đáp",
        "context_hint": "VB phúc đáp",
        "forms": [
            {"type": "Công văn", "when": "Mặc định cho phúc đáp"},
            {"type": "Báo cáo", "when": "Khi VB nguồn yêu cầu báo cáo phúc đáp"},
        ],
    },
    {
        "pattern": r"đề\s+xuất\s+(?!chương\s+trình)",
        "context_hint": "VB đề xuất (chưa rõ cấp ban hành)",
        "forms": [
            {"type": "Tờ trình", "when": "Đề xuất lên cấp trên có thẩm quyền quyết định"},
            {"type": "Công văn", "when": "Đề xuất ngang cấp / phối hợp"},
            {"type": "Báo cáo", "when": "Đề xuất kèm tổng kết/đánh giá tình hình"},
        ],
    },
]


def _get_classify_rules() -> list[tuple[str, str]]:
    """Load classify rules từ YAML (cache→bundled), fallback hardcode nếu thiếu."""
    data = load_rules("loai-vb")
    if not data or not data.get("classify_rules"):
        return _CLASSIFY_FALLBACK
    return [(r["pattern"], r["type"]) for r in data["classify_rules"]
            if r.get("pattern") and r.get("type")]


def _get_ambiguous_patterns() -> list[dict]:
    """Load ambiguous forms từ YAML, fallback hardcode nếu thiếu."""
    data = load_rules("loai-vb")
    if not data or not data.get("ambiguous_forms"):
        return _AMBIGUOUS_FALLBACK
    return data["ambiguous_forms"]


# Backward-compat aliases (code cũ có thể import 2 biến này)
CLASSIFY_RULES = _get_classify_rules()
AMBIGUOUS_FORM_PATTERNS = _get_ambiguous_patterns()


# ---- MCP server ----

mcp = FastMCP("vbhc")


@mcp.tool()
def vbhc_classify(description: str) -> dict:
    """Classify a Vietnamese administrative document type from a free-text description.

    Args:
        description: User's description of what they want to draft (1-3 sentences, Vietnamese).

    Returns:
        dict with:
            - matches: list of {type, confidence} for matched VB types
            - suggestion: top match (or "Không xác định" if no match)
            - need_clarification: bool (true if 0 or 2+ matches)
    """
    desc_lower = description.lower()
    matches = []
    # Gọi helper mỗi call → sau khi vbhc_sync_knowledge (Phase 2) clear_cache(),
    # lần classify tiếp theo sẽ dùng rules YAML mới nhất.
    for pattern, label in _get_classify_rules():
        if re.search(pattern, desc_lower):
            matches.append({"type": label, "confidence": "high"})

    # Detect ambiguous-form context (góp ý / phúc đáp / đề xuất...)
    ambiguous = []
    for rule in _get_ambiguous_patterns():
        if re.search(rule["pattern"], desc_lower):
            ambiguous.append({
                "context": rule["context_hint"],
                "forms": rule["forms"],
            })

    if not matches and not ambiguous:
        return {
            "matches": [],
            "suggestion": "Không xác định",
            "need_clarification": True,
            "hint": "Hỏi user mục đích cụ thể: thông báo / yêu cầu / trình lên / quyết định ban hành / báo cáo / xin ý kiến",
        }

    result = {
        "matches": matches,
        "suggestion": matches[0]["type"] if matches else None,
        "need_clarification": len(matches) > 1 or bool(ambiguous),
    }

    if ambiguous:
        result["ambiguous_forms"] = ambiguous
        result["clarification_required"] = (
            "Đây là VB phản hồi/góp ý/đề xuất — dạng VB phụ thuộc vào YÊU CẦU của VB nguồn. "
            "HỎI user chọn dạng (Báo cáo / Công văn / Tờ trình) trước khi soạn. "
            "Nếu user có VB nguồn → đọc xem nguồn yêu cầu dạng nào."
        )

    return result


@mcp.tool()
def vbhc_create_workfolder(
    description: str,
    parent_dir: str,
    custom_slug: str | None = None,
) -> dict:
    """Create a new work folder following the <NNNN>-<slug>/ convention.

    Args:
        description: Free-text description (used to generate slug if custom_slug not given).
        parent_dir: Parent directory path. Will be created if missing.
        custom_slug: Optional explicit slug (without NNNN-).

    Returns:
        dict with: folder_path, structure (list of created paths)
    """
    parent = Path(parent_dir).resolve()
    nnnn = next_folder_number(parent)
    slug = slugify_vn(custom_slug or description)[:60] or "ho-so-moi"
    folder = parent / f"{nnnn}-{slug}"

    if folder.exists():
        return {"error": f"Folder already exists: {folder}", "folder_path": str(folder)}

    folder.mkdir(parents=True)
    ky_thuat = folder / "0-ky-thuat"
    ky_thuat.mkdir()
    (folder / "1-tham-chieu").mkdir()

    # Copy templates → 0-ky-thuat/
    tpl_req = SKILL_DIR / "resources" / "templates" / "1-yeu-cau.md.tpl"
    tpl_yaml = SKILL_DIR / "resources" / "templates" / "2-du-lieu.yaml.tpl"

    yc = ky_thuat / "1-yeu-cau.md"
    if tpl_req.exists():
        yc.write_text(tpl_req.read_text(encoding="utf-8"), encoding="utf-8")
    else:
        yc.write_text("# Yêu cầu công việc\n", encoding="utf-8")

    dl = ky_thuat / "2-du-lieu.yaml"
    if tpl_yaml.exists():
        dl.write_text(tpl_yaml.read_text(encoding="utf-8"), encoding="utf-8")
    else:
        dl.write_text("# Dữ liệu\n", encoding="utf-8")

    return {
        "folder_path": str(folder),
        "structure": [
            str(ky_thuat / "1-yeu-cau.md"),
            str(ky_thuat / "2-du-lieu.yaml"),
            str(folder / "1-tham-chieu/"),
        ],
        "next_step": "Phỏng vấn user để điền 1-yeu-cau.md và 2-du-lieu.yaml",
    }


@mcp.tool()
def vbhc_reorganize(
    source_folder: str,
    custom_slug: str | None = None,
    parent_dir: str | None = None,
) -> dict:
    """Reorganize a messy folder into the <NNNN>-<slug>/ convention.

    Moves all files from source_folder into <new>/3-tham-chieu/, creates
    metadata templates, and tries to remove the old folder.

    Args:
        source_folder: Path to messy folder.
        custom_slug: Optional slug for new folder.
        parent_dir: Where to put new folder (default: same parent as source).

    Returns:
        dict with: new_folder_path, files_moved, old_folder_status
    """
    src = Path(source_folder).resolve()
    if not src.is_dir():
        return {"error": f"Source not a directory: {src}"}

    parent = Path(parent_dir).resolve() if parent_dir else src.parent
    parent.mkdir(parents=True, exist_ok=True)

    nnnn = next_folder_number(parent)
    slug = slugify_vn(custom_slug or src.name)[:60] or "ho-so-moi"
    new_folder = parent / f"{nnnn}-{slug}"

    if new_folder.exists():
        return {"error": f"Target exists: {new_folder}"}

    new_folder.mkdir(parents=True)
    ky_thuat = new_folder / "0-ky-thuat"
    tham_chieu = new_folder / "1-tham-chieu"
    ky_thuat.mkdir()
    tham_chieu.mkdir()

    files = [f for f in src.iterdir() if f.is_file()]
    moved_names = []
    for f in files:
        shutil.move(str(f), str(tham_chieu / f.name))
        moved_names.append(f.name)

    # Templates → 0-ky-thuat/
    tpl_req = SKILL_DIR / "resources" / "templates" / "1-yeu-cau.md.tpl"
    tpl_yaml = SKILL_DIR / "resources" / "templates" / "2-du-lieu.yaml.tpl"
    if tpl_req.exists():
        (ky_thuat / "1-yeu-cau.md").write_text(tpl_req.read_text(encoding="utf-8"), encoding="utf-8")
    if tpl_yaml.exists():
        (ky_thuat / "2-du-lieu.yaml").write_text(tpl_yaml.read_text(encoding="utf-8"), encoding="utf-8")

    # Try remove old
    try:
        src.rmdir()
        old_status = "deleted"
    except OSError as e:
        old_status = f"NOT deleted (locked or non-empty): {e}"

    return {
        "new_folder_path": str(new_folder),
        "files_moved": moved_names,
        "files_count": len(moved_names),
        "old_folder_status": old_status,
        "next_step": "Đọc các file trong 3-tham-chieu/, phỏng vấn user, fill 1-yeu-cau.md + 2-du-lieu.yaml",
    }


@mcp.tool()
def vbhc_fill_template(
    template_path: str,
    output_path: str,
    cell_ops: list[dict] | None = None,
    paragraph_ops: list[dict] | None = None,
    replace_ops: list[dict] | None = None,
) -> dict:
    """Fill a .docx template by editing specific cells, paragraphs, or doing global replace.

    Handles the case where mcp__word__search_and_replace fails on text inside
    table cells (common after .doc → .docx conversion).

    Args:
        template_path: Source .docx — chấp nhận 2 dạng:
            • Path tuyệt đối/tương đối (chứa / hoặc \\) → dùng nguyên path
            • Slug (vd "bao-cao", "cong-van") → tự resolve từ ~/.vbhc/cache/
              fallback bundled SKILL_DIR/resources/templates/. Tự auto-pull
              từ cloud KB nếu cache thiếu và đã config.
        output_path: Where to save the filled .docx.
        cell_ops: List of {table_idx, row_idx, col_idx, text} operations.
        paragraph_ops: List of {paragraph_idx, text} operations.
        replace_ops: List of {old, new} text replacements (applied to all paragraphs and cells).

    Returns:
        dict with: output_path, ops_applied, warnings
    """
    src = _resolve_template_path(template_path)
    dst = Path(output_path).expanduser().resolve()

    if not src.is_file():
        return {"error": f"Template not found: {src} (resolved from {template_path!r})"}

    if src != dst:
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(src, dst)

    doc = Document(str(dst))
    ops_applied = []
    warnings = []

    for op in (cell_ops or []):
        try:
            cell = doc.tables[op["table_idx"]].rows[op["row_idx"]].cells[op["col_idx"]]
            cell_set_text(cell, op["text"])
            ops_applied.append(f"cell T{op['table_idx']}/R{op['row_idx']}/C{op['col_idx']}")
        except (IndexError, KeyError) as e:
            warnings.append(f"cell op failed: {op} ({e})")

    for op in (paragraph_ops or []):
        try:
            replace_paragraph_text(doc.paragraphs[op["paragraph_idx"]], op["text"])
            ops_applied.append(f"paragraph P{op['paragraph_idx']}")
        except (IndexError, KeyError) as e:
            warnings.append(f"paragraph op failed: {op} ({e})")

    for op in (replace_ops or []):
        n = search_and_replace_doc(doc, op["old"], op["new"])
        if n:
            ops_applied.append(f"replace '{op['old'][:30]}' ×{n}")
        else:
            warnings.append(f"replace not found: {op['old'][:60]}")

    doc.save(str(dst))

    # Auto-validate ND30 sau khi save (NĐ30 quy định BẮT BUỘC mỗi VB ban hành phải
    # tuân thủ 9 thành phần thể thức). Nếu có ✗ → đưa vào warnings để AI thấy.
    nd30_check = _validate_nd30(dst)
    if nd30_check["fail_count"] > 0:
        warnings.append(
            f"ND30: {nd30_check['fail_count']} mục FAIL — file CHƯA tuân thủ. "
            f"Xem nd30_validation."
        )

    return {
        "output_path": str(dst),
        "ops_applied": ops_applied,
        "warnings": warnings,
        "nd30_validation": nd30_check,
    }


def _validate_nd30(docx_path: Path) -> dict:
    """Helper: chạy validate_thethuc 9 mục, trả counts + summary list. Dùng nội bộ
    bởi vbhc_fill_template (auto-hook) và vbhc_validate (public tool)."""
    import validate_thethuc as vt
    doc = Document(str(docx_path))
    text = vt.collect_all_text(doc)
    items = [
        ("1. Quốc hiệu + Tiêu ngữ",   vt.check_quoc_hieu(text)),
        ("2. Tên cơ quan ban hành",   vt.check_co_quan(text)),
        ("3. Số/ký hiệu",             vt.check_so_van_ban(text)),
        ("4. Tên loại + Trích yếu",   vt.check_ten_loai(text)),
        ("5. Nội dung",               vt.check_noi_dung(text)),
        ("6. Người ký",               vt.check_nguoi_ky(text)),
        ("7. Dấu/chữ ký số",          vt.check_dau()),
        ("8. Nơi nhận + Lưu",         vt.check_noi_nhan(text)),
        ("9. Phụ lục",                vt.check_phu_luc(text)),
    ]
    summary = [
        {"label": label, "status": status, "detail": detail}
        for label, (status, detail) in items
    ]
    ok = sum(1 for it in summary if it["status"] == "✓")
    warn = sum(1 for it in summary if it["status"] == "⚠")
    fail = sum(1 for it in summary if it["status"] == "✗")
    return {"summary": summary, "ok_count": ok, "warn_count": warn, "fail_count": fail}


@mcp.tool()
def vbhc_validate(docx_path: str) -> dict:
    """Run thể thức checklist (9 components per NĐ 30/2020) on a .docx file.

    Args:
        docx_path: Path to .docx file to check.

    Returns:
        dict with: passed, warned, failed counts + per-item results.
    """
    path = Path(docx_path).resolve()
    if not path.is_file():
        return {"error": f"File not found: {path}"}

    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    text = "\n".join(parts)

    is_bm = bool(re.search(r"PHIẾU\s+(GHI\s+Ý\s+KIẾN|BIỂU\s+QUYẾT|THẨM\s+ĐỊNH|LẤY\s+Ý\s+KIẾN)", text))

    placeholder_re = re.compile(r"\?\?\?|<[^>]+>|\[placeholder\]", re.IGNORECASE)

    items = []

    # 1
    if "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in text and "Hạnh phúc" in text:
        items.append(("Quốc hiệu + Tiêu ngữ", "ok", "OK"))
    else:
        items.append(("Quốc hiệu + Tiêu ngữ", "fail", "Thiếu hoặc sai"))

    # 2
    upper_lines = [ln for ln in text.split("\n") if ln.strip() and ln == ln.upper() and len(ln) > 4]
    if any(re.search(r"(UBND|ỦY BAN|BỘ|SỞ|HĐND|VĂN PHÒNG|TỔNG CỤC|CỤC|VIỆN|BAN)", ln) for ln in upper_lines):
        items.append(("Tên cơ quan", "ok", "OK"))
    else:
        items.append(("Tên cơ quan", "warn", "Không phát hiện được"))

    # 3
    if is_bm:
        items.append(("Số/ký hiệu", "ok", "Không cần (biểu mẫu nội bộ)"))
    else:
        m = re.search(r"Số:\s*([0-9]*\s*/[A-ZĐa-z&\-_]+)", text)
        if m:
            sotxt = m.group(1).strip()
            if sotxt.startswith("/"):
                items.append(("Số/ký hiệu", "warn", f"Trống: {m.group(0).strip()}"))
            else:
                items.append(("Số/ký hiệu", "ok", m.group(0).strip()))
        else:
            items.append(("Số/ký hiệu", "fail", "Thiếu 'Số:'"))

    # 4
    keywords = ["QUYẾT ĐỊNH", "NGHỊ QUYẾT", "BÁO CÁO", "TỜ TRÌNH", "THÔNG BÁO", "KẾ HOẠCH",
                "CHỈ THỊ", "BIÊN BẢN", "GIẤY MỜI", "PHIẾU GHI Ý KIẾN", "PHIẾU BIỂU QUYẾT",
                "HƯỚNG DẪN", "KẾT LUẬN"]
    found = [k for k in keywords if k in text]
    if found:
        items.append(("Tên loại + Trích yếu", "ok", found[0]))
    elif "V/v" in text:
        items.append(("Tên loại + Trích yếu", "ok", "Công văn (V/v)"))
    else:
        items.append(("Tên loại + Trích yếu", "warn", "Không xác định"))

    # 5
    ph = placeholder_re.findall(text)
    if ph:
        items.append(("Nội dung", "fail", f"Còn placeholder: {ph[:3]}"))
    else:
        items.append(("Nội dung", "ok", "OK"))

    # 6
    chuc_vu = re.search(r"(KT\.|TL\.|TUQ\.)?\s*(GIÁM ĐỐC|CHỦ TỊCH|PHÓ CHỦ TỊCH|CHÁNH VĂN PHÒNG|"
                        r"PHÓ CHÁNH VĂN PHÒNG|TRƯỞNG PHÒNG|PHÓ GIÁM ĐỐC|BỘ TRƯỞNG|THỨ TRƯỞNG|"
                        r"VỤ TRƯỞNG|CỤC TRƯỞNG|VIỆN TRƯỞNG)", text)
    if chuc_vu:
        items.append(("Người ký", "ok", chuc_vu.group(0).strip()))
    else:
        items.append(("Người ký", "warn", "Không phát hiện chức vụ in hoa"))

    # 7
    items.append(("Dấu / chữ ký số", "warn", "Kiểm tra thủ công"))

    # 8
    if is_bm:
        items.append(("Nơi nhận + Lưu", "ok", "Không cần (biểu mẫu nội bộ)"))
    elif "Nơi nhận:" in text:
        if "Lưu:" in text:
            items.append(("Nơi nhận + Lưu", "ok", "OK"))
        else:
            items.append(("Nơi nhận + Lưu", "warn", "Thiếu 'Lưu:'"))
    else:
        items.append(("Nơi nhận + Lưu", "fail", "Thiếu"))

    # 9
    has_kem = bool(re.search(r"kèm theo|đính kèm", text, re.IGNORECASE))
    has_phu_luc = bool(re.search(r"PHỤ LỤC\s+[IVX]+", text))
    if has_kem and not has_phu_luc:
        items.append(("Phụ lục", "warn", "Có nhắc 'kèm theo' nhưng không tìm thấy phụ lục"))
    elif has_phu_luc:
        items.append(("Phụ lục", "ok", "OK"))
    else:
        items.append(("Phụ lục", "ok", "Không có (không bắt buộc)"))

    passed = sum(1 for _, s, _ in items if s == "ok")
    warned = sum(1 for _, s, _ in items if s == "warn")
    failed = sum(1 for _, s, _ in items if s == "fail")

    return {
        "passed": passed,
        "warned": warned,
        "failed": failed,
        "total": len(items),
        "items": [{"name": n, "status": s, "detail": d} for n, s, d in items],
        "is_bieu_mau_noi_bo": is_bm,
    }


@mcp.tool()
def vbhc_aggregate_survey(
    xlsx_path: str,
    role_col: int | None = None,
    max_comments_per_col: int = 50,
) -> dict:
    """Aggregate a Google Forms-style .xlsx survey: stats + non-trivial comments.

    Use case: cấp dưới đã thu thập ý kiến qua Google Forms; cần tổng hợp trước khi
    soạn báo cáo góp ý.

    Args:
        xlsx_path: Path to .xlsx file (e.g. Google Forms export).
        role_col: 1-indexed column for respondent role. None = auto-detect.
        max_comments_per_col: Max comments to return per text column (truncate older).

    Returns:
        dict with: total_responses, demographics, stats (per numeric col), comments (per text col)
    """
    if not HAS_OPENPYXL:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    import re as _re
    path = Path(xlsx_path).resolve()
    if not path.is_file():
        return {"error": f"File not found: {path}"}

    TRIVIAL = _re.compile(
        r"^(không|không\s+có\s+ý\s+kiến|không\s+có|không\s+có\s+(đề|đè)\s+xuất|"
        r"k|kh|o|/|-+|không\s+có\s+gì|nothing|no|none|x|\.|chưa\s+có|chưa|"
        r"đồng\s+ý|nhất\s+trí|không\s+(có\s+)?(ý\s+kiến|gì|đề\s+xuất|bổ\s+sung).*)$",
        _re.IGNORECASE,
    )
    COMMENT_KW = ["góp ý", "đề xuất", "ý kiến", "bình luận", "comment",
                  "feedback", "kiến nghị", "khuyến nghị", "nhận xét"]

    wb = openpyxl.load_workbook(str(path), data_only=True)
    ws = wb.active

    headers = []
    for c in range(1, ws.max_column + 1):
        v = ws.cell(row=1, column=c).value
        headers.append(str(v).strip() if v else "")

    if role_col is None:
        for i, h in enumerate(headers, 1):
            hl = h.lower()
            if "vai trò" in hl or "role" in hl or "đối tượng" in hl:
                role_col = i
                break

    rows = []
    for r in range(2, ws.max_row + 1):
        row = [ws.cell(row=r, column=c).value for c in range(1, ws.max_column + 1)]
        if all(v is None or (isinstance(v, str) and not v.strip()) for v in row):
            continue
        rows.append(row)

    n = len(rows)
    result: dict[str, Any] = {"total_responses": n, "headers": headers}

    if role_col:
        roles_count: dict[str, int] = {}
        for r in rows:
            v = r[role_col-1]
            if v:
                roles_count[str(v).strip()] = roles_count.get(str(v).strip(), 0) + 1
        result["demographics"] = [
            {"role": k, "count": v, "percent": round(100*v/n, 1)}
            for k, v in sorted(roles_count.items(), key=lambda x: -x[1])
        ]

    # Stats per numeric col
    stats = []
    for ci, h in enumerate(headers, 1):
        vals = []
        for r in rows:
            v = r[ci-1]
            if v is None:
                continue
            try:
                vals.append(float(v))
            except (TypeError, ValueError):
                pass
        if not vals:
            continue
        stats.append({
            "col": ci,
            "header": h[:200],
            "n": len(vals),
            "avg": round(sum(vals)/len(vals), 3),
            "min": min(vals),
            "max": max(vals),
        })
    result["stats"] = stats

    # Comments per text col
    comments_out = []
    for ci, h in enumerate(headers, 1):
        hl = h.lower()
        vals = [str(r[ci-1]) for r in rows if r[ci-1] is not None and isinstance(r[ci-1], str)]
        if not vals or len(vals) < 5:
            continue
        unique_count = len({v.strip().lower() for v in vals})
        if unique_count < 5:
            continue
        is_comment = any(kw in hl for kw in COMMENT_KW) or (
            sum(len(v) for v in vals) / len(vals) > 50
        )
        if not is_comment:
            continue
        items = []
        for r in rows:
            v = r[ci-1]
            if v is None:
                continue
            t = str(v).strip()
            if not t or len(t) < 4 or TRIVIAL.fullmatch(t):
                continue
            role = str(r[role_col-1]).strip() if role_col else ""
            items.append({"role": role, "text": t[:500]})
        comments_out.append({
            "col": ci,
            "header": h[:200],
            "non_trivial_count": len(items),
            "items": items[:max_comments_per_col],
        })
    result["comments"] = comments_out

    return result


@mcp.tool()
def vbhc_regenerate_check(work_folder: str, update: bool = False) -> dict:
    """Detect new/changed files in 1-tham-chieu/ since last regenerate.

    Use case: user nói "soạn lại VB này" sau khi quăng thêm file vào folder.
    Trước khi gen lại, AI cần biết có file mới hoặc file cũ thay đổi không.

    Args:
        work_folder: Path to <NNNN>-<...> work folder
        update: If True, update file-manifest.yaml after reporting

    Returns:
        dict with new/changed/removed/unchanged file lists + manifest path
    """
    import time
    work = Path(work_folder).resolve()
    if not work.is_dir():
        return {"error": f"Not a directory: {work}"}

    tham_chieu = work / "1-tham-chieu"
    ky_thuat = work / "0-ky-thuat"
    manifest_path = ky_thuat / "file-manifest.yaml"

    if not tham_chieu.is_dir():
        return {"error": f"1-tham-chieu/ not found in {work}"}
    ky_thuat.mkdir(exist_ok=True)

    # Scan current
    current = []
    for f in tham_chieu.iterdir():
        if not f.is_file() or f.name.startswith("~$"):
            continue
        st = f.stat()
        current.append({"name": f.name, "size": st.st_size, "mtime": int(st.st_mtime)})
    current.sort(key=lambda x: x["name"])

    # Read previous manifest (simple YAML parsing)
    previous = []
    if manifest_path.is_file():
        cur_item = None
        for line in manifest_path.read_text(encoding="utf-8").splitlines():
            s = line.strip()
            if not s or s.startswith("#"):
                continue
            if s.startswith("- name:"):
                if cur_item:
                    previous.append(cur_item)
                cur_item = {"name": s.split(":", 1)[1].strip().strip('"')}
            elif cur_item is not None and ":" in s:
                k, v = s.split(":", 1)
                k = k.strip()
                v = v.strip().strip('"')
                try:
                    cur_item[k] = int(v)
                except ValueError:
                    cur_item[k] = v
        if cur_item:
            previous.append(cur_item)

    # Diff
    prev_by_name = {f["name"]: f for f in previous}
    new = [f for f in current if f["name"] not in prev_by_name]
    changed = []
    unchanged = []
    for f in current:
        if f["name"] in prev_by_name:
            p = prev_by_name[f["name"]]
            if f["size"] != p.get("size") or f["mtime"] != p.get("mtime"):
                changed.append({**f, "previous_size": p.get("size")})
            else:
                unchanged.append(f)
    cur_names = {f["name"] for f in current}
    removed = [f for f in previous if f["name"] not in cur_names]

    if update:
        today = time.strftime("%Y-%m-%d")
        lines = ["# File manifest — auto-generated by vbhc_regenerate_check",
                 "files:"]
        for f in current:
            lines.append(f'  - name: "{f["name"]}"')
            lines.append(f'    size: {f["size"]}')
            lines.append(f'    mtime: {f["mtime"]}')
            lines.append(f'    last_read: "{today}"')
        manifest_path.write_text("\n".join(lines), encoding="utf-8")

    return {
        "work_folder": str(work),
        "manifest_path": str(manifest_path),
        "current_count": len(current),
        "previous_count": len(previous),
        "new": new,
        "changed": changed,
        "removed": removed,
        "unchanged_count": len(unchanged),
        "needs_re_read": bool(new or changed),
        "first_run": not previous,
    }


@mcp.tool()
def vbhc_load_org_config(filename: str = "05-thong-tin-co-quan.yaml") -> dict:
    """Load a config YAML from the ORG tier (cơ quan share).

    The ORG dir is resolved from $VBHC_ORG_DIR (default ~/.vbhc/org/).
    Common files:
      - 05-thong-tin-co-quan.yaml  → tên cơ quan, người ký, phòng, viết tắt
      - phan-cong-nhiem-vu.yaml    → danh sách phòng + chức năng nhiệm vụ
      - can-cu-phap-ly-mau.yaml    → danh sách VB pháp lý mẫu (verify hiệu lực)

    Args:
        filename: file name (relative to ORG dir).

    Returns:
        dict with: org_dir, file_path, exists, content (raw text), parsed (yaml or null)
    """
    org_dir = ORG_DIR
    f = (org_dir / filename).resolve()

    result = {
        "org_dir": str(org_dir),
        "file_path": str(f),
        "exists": f.is_file(),
        "content": None,
        "parsed": None,
        "hint": None,
    }

    if not org_dir.is_dir():
        result["hint"] = (
            f"ORG dir chưa tồn tại: {org_dir}. "
            f"Tạo folder + copy template từ {SKILL_DIR / 'tri-thuc-template'}/. "
            f"Hoặc set env var VBHC_ORG_DIR trỏ đến folder khác."
        )
        return result

    if not f.is_file():
        result["hint"] = (
            f"File chưa có trong ORG dir. Cơ quan cần tạo {filename} trong {org_dir}. "
            f"Xem template tại {SKILL_DIR / 'tri-thuc-template' / filename} (nếu có)."
        )
        return result

    text = f.read_text(encoding="utf-8")
    result["content"] = text
    try:
        import yaml
        result["parsed"] = yaml.safe_load(text)
    except ImportError:
        result["hint"] = "PyYAML not installed — return raw content only. Run: pip install pyyaml"
    except Exception as e:
        result["hint"] = f"YAML parse error: {e}"

    return result


@mcp.tool()
def vbhc_suggest_noi_nhan(
    vb_purpose: str,
    vb_type: str = "",
    user_provided: list[str] | None = None,
) -> dict:
    """Gợi ý nơi nhận dựa trên loại VB + mục đích + phân công nhiệm vụ của cơ quan.

    AI gọi tool này TRƯỚC khi đưa nơi nhận vào VB. Tool đọc:
      - phan-cong-nhiem-vu.yaml (ORG tier) → danh sách cơ quan/đơn vị + chức năng
      - 05-thong-tin-co-quan.yaml (ORG tier) → noi_nhan_phobien

    Args:
        vb_purpose: Mục đích VB (vd: "góp ý dự thảo Thông tư của Bộ X")
        vb_type: Loại VB (Báo cáo / Công văn / Tờ trình / ...)
        user_provided: Danh sách user đã nói (sẽ keep + sanitize)

    Returns:
        dict với: required (bắt buộc), suggested (đề xuất), to_ask_user (cần xác nhận),
                  warning (nếu thiếu phân công NV → đề nghị user cung cấp)
    """
    org_cfg = vbhc_load_org_config("05-thong-tin-co-quan.yaml")
    pc_nv_cfg = vbhc_load_org_config("phan-cong-nhiem-vu.yaml")

    result: dict[str, Any] = {
        "required": [],
        "suggested": [],
        "to_ask_user": [],
        "warning": None,
        "user_provided_sanitized": [],
    }

    # Sanitize user_provided: bỏ "(...)" mục đích trong ngoặc
    if user_provided:
        for item in user_provided:
            cleaned = re.sub(r"\s*\([^)]*\)\s*\.?\s*$", "", item).rstrip(" .")
            result["user_provided_sanitized"].append(cleaned)

    # Required: luôn có "Lưu: VT, <viet_tat_phong>." → AI điền sau khi user chọn phòng
    result["required"].append("Lưu: VT, <viet_tat_phong soạn thảo>.")

    # Suggested theo logic loại VB
    purpose_low = vb_purpose.lower()
    if vb_type == "Báo cáo" or "báo cáo" in purpose_low:
        # Báo cáo → cấp trên trực tiếp
        result["suggested"].append("Cơ quan ban hành VB nguồn (vd: Bộ ngành chủ quản)")
        result["suggested"].append("UBND tỉnh (nếu cần báo cáo cấp tỉnh)")
    if vb_type == "Tờ trình":
        result["suggested"].append("Cấp có thẩm quyền quyết định (UBND tỉnh / Chủ tịch / Bộ trưởng)")
    if "góp ý" in purpose_low or "phản hồi" in purpose_low:
        result["suggested"].append("Cơ quan đã gửi VB xin ý kiến (BẮT BUỘC)")

    # Phân công NV available?
    if not pc_nv_cfg["exists"]:
        result["warning"] = (
            "⚠ Chưa có phan-cong-nhiem-vu.yaml trong ORG dir. "
            "AI ĐỀ NGHỊ user cung cấp phân công nhiệm vụ của cơ quan để gợi ý nơi nhận chính xác hơn "
            "(tránh gửi thừa/thiếu). Tạm thời chỉ gợi ý từ logic loại VB + danh mục phổ biến."
        )
    else:
        # Đọc phân công, gợi ý cơ quan/đơn vị có chức năng liên quan
        parsed = pc_nv_cfg.get("parsed") or {}
        donvi = parsed.get("don_vi", []) if isinstance(parsed, dict) else []
        for d in donvi:
            if not isinstance(d, dict):
                continue
            chuc_nang = " ".join(d.get("chuc_nang", []) if isinstance(d.get("chuc_nang"), list) else [str(d.get("chuc_nang", ""))])
            if any(kw in (chuc_nang + " " + d.get("ten", "")).lower() for kw in purpose_low.split() if len(kw) > 3):
                result["to_ask_user"].append({
                    "ten": d.get("ten", ""),
                    "ly_do": f"Có chức năng liên quan: {chuc_nang[:120]}",
                })

    result["instruction"] = (
        "AI: HỎI user xác nhận từng nhóm trước khi đưa vào VB. "
        "KHÔNG ghi mục đích trong ngoặc đơn (vd: '(để báo cáo)') — NĐ 30 không quy định."
    )
    return result


@mcp.tool()
def vbhc_learn_template(file_path: str) -> dict:
    """Đọc 1 file Word VBHC mẫu của user, phân tích thể thức theo NĐ30, trả về
    spec + report (Markdown) để user xem trước khi quyết định lưu làm template.

    AI gọi tool này khi user nói: "học mẫu này", "phân tích mẫu", "kiểm tra thể
    thức file X", "soi mẫu này", hoặc khi muốn so sánh 1 file với chuẩn ND30.

    Output `report_md` là báo cáo human-readable — AI nên hiển thị nguyên văn cho
    user thay vì tự diễn giải, để user thấy điểm sai ở đâu, đề xuất sửa ra sao.

    Args:
        file_path: đường dẫn tuyệt đối tới file .docx (mẫu user)

    Returns:
        dict gồm: file, loai_vb (slug), spec (chi tiết), issues (cần sửa),
                  validation (9 mục), report_md
    """
    p = Path(file_path).expanduser().resolve()
    if not p.is_file():
        return {"error": f"file not found: {p}"}
    if p.suffix.lower() != ".docx":
        return {"error": f"file phải có đuôi .docx, không phải {p.suffix}"}

    import learn_template as lt
    spec, validation, issues = lt.learn(p)
    return {
        "file": p.name,
        "loai_vb": spec["loai_vb"],
        "spec": spec,
        "issues": issues,
        "validation": [
            {"label": label, "status": status, "detail": detail}
            for label, (status, detail) in validation
        ],
        "report_md": lt.build_report(spec, validation, issues),
    }


@mcp.tool()
def vbhc_update_template(
    source_file: str,
    target_loai_vb: str,
    confirmed: bool = False,
) -> dict:
    """Lưu 1 file Word đã được duyệt làm TEMPLATE chuẩn cho 1 loại VB,
    ghi vào LOCAL CACHE (~/.vbhc/cache/templates/<target_loai_vb>.docx).

    AI gọi tool này khi user nói: "cập nhật mẫu", "lưu file này thành template",
    "thay template loại X bằng file này", thường sau khi đã chạy
    vbhc_learn_template và user duyệt nội dung.

    Workflow chuẩn (LUÔN qua 2 bước):
      1. confirmed=False → tool trả preview + cảnh báo (file đích đã có hay chưa,
         còn lỗi ND30 cần sửa không). AI hiển thị cho user.
      2. User OK → confirmed=True → tool copy file vào ~/.vbhc/cache/templates/.

    Sau khi ghi cache, template được dùng NGAY bởi vbhc_fill_template (qua
    _resolve_template_path → cache trước, bundled sau). Để chia sẻ cho máy
    khác, dùng vbhc_publish_template (Phase 4) để push lên cloud KB Hub.

    AN TOÀN: nếu source FAIL bất kỳ mục ND30 nào → REFUSE save, trả lại issues.

    Args:
        source_file: đường dẫn file .docx nguồn (do user cung cấp)
        target_loai_vb: slug loại VB (vd: "bao-cao", "cong-van", "phieu-ghi-y-kien")
        confirmed: True khi user xác nhận ghi đè (lần gọi thứ 2)

    Returns:
        dict gồm: source, target, exists, confirmed, applied, validation_fails,
                  issues_to_fix_first, preview_message hoặc message
    """
    src = Path(source_file).expanduser().resolve()
    if not src.is_file():
        return {"error": f"file source not found: {src}"}
    if src.suffix.lower() != ".docx":
        return {"error": f"file phải có đuôi .docx, không phải {src.suffix}"}

    slug = re.sub(r"[^a-z0-9\-]", "-", target_loai_vb.lower()).strip("-")
    if not slug:
        return {"error": f"target_loai_vb không hợp lệ: '{target_loai_vb}'"}

    target = kc.cache_path_for("templates", f"{slug}.docx")

    # Validate ND30 trước khi cho phép save làm template chuẩn
    import learn_template as lt
    spec, validation, issues = lt.learn(src)
    fail_count = sum(1 for _, (status, _) in validation if status == "✗")

    result: dict[str, Any] = {
        "source": str(src),
        "target": str(target),
        "loai_vb_detected": spec["loai_vb"],
        "loai_vb_target": slug,
        "exists": target.is_file(),
        "confirmed": confirmed,
        "applied": False,
        "validation_fails": fail_count,
        "issues_to_fix_first": [iss for iss in issues if iss["level"] == "fix"],
    }

    if fail_count > 0:
        result["error"] = (
            f"File source FAIL {fail_count} mục thể thức NĐ30. "
            f"PHẢI fix trước khi lưu làm template. Xem issues_to_fix_first."
        )
        return result

    if not confirmed:
        warn = "⚠ File đích đã tồn tại trong cache — sẽ GHI ĐÈ. " if target.is_file() else ""
        result["preview_message"] = (
            f"Sẽ copy '{src.name}' → cache '{target}'. {warn}"
            f"Gọi lại với confirmed=True để thực hiện. "
            f"Để chia sẻ lên cloud cho máy khác: dùng vbhc_publish_template "
            f"(Phase 4) sau khi cache ghi xong."
        )
        return result

    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(src), str(target))
    result["applied"] = True
    result["message"] = (
        f"✓ Đã lưu template chuẩn vào cache: {target}. "
        f"vbhc_fill_template dùng được ngay. Để chia sẻ cho máy khác: "
        f"vbhc_publish_template('{slug}') (Phase 4)."
    )
    return result


# =====================================================================
# Cloud sync tools (Phase 2)
# =====================================================================

@mcp.tool()
def vbhc_sync_knowledge(force: bool = False, only: str | None = None) -> dict:
    """Pull templates + rules + code bundle từ cloud KB Hub về ~/.vbhc/cache/.

    Args:
        force: pull cả khi cache còn fresh (TTL chưa hết).
        only: chỉ pull 1 kind: "manifest", "templates", "rules", "code".
              None = sync tất cả (recommended).

    Returns:
        dict với counts: templates/rules/code/errors. Sau khi sync, các tool
        khác (vbhc_classify, vbhc_validate, ...) sẽ dùng version mới ở lần
        gọi tiếp theo (rules_loader.clear_cache() được gọi tự động).
    """
    if not kc.is_configured():
        return {
            "error": "Chưa bootstrap — chạy `python -m mcp.bootstrap` trước, "
                     "hoặc tạo ~/.vbhc/config.yaml với cloud_url + api_key."
        }

    try:
        if only == "manifest":
            info = kc.sync_manifest()
            from rules_loader import clear_cache
            clear_cache()
            return {
                "ok": True,
                "manifest": {
                    "status": info["status"],
                    "templates": list((info["manifest"].get("templates") or {}).keys()),
                    "rules": list((info["manifest"].get("rules") or {}).keys()),
                    "code_version": (info["manifest"].get("code") or {}).get("version"),
                },
            }
        summary = kc.sync_all()
    except kc.KBError as e:
        return {"error": str(e)}

    # Buộc các rule YAML reload từ disk sau khi sync
    from rules_loader import clear_cache
    clear_cache()

    return {"ok": True, **summary}


@mcp.tool()
def vbhc_publish_template(slug: str, confirmed: bool = False) -> dict:
    """Đẩy 1 template từ local cache (`~/.vbhc/cache/templates/<slug>.docx`)
    lên cloud KB Hub để chia sẻ cho mọi máy khác trong cơ quan.

    Yêu cầu: API key trong ~/.vbhc/config.yaml phải có scope `admin`.

    Workflow chuẩn (LUÔN qua 2 bước):
      1. confirmed=False → tool kiểm tra cache có file không, in preview
         (slug, size, sha256 local) + cảnh báo nếu cache là bundled file
         (chưa qua vbhc_update_template). AI hiển thị cho user.
      2. User OK → confirmed=True → tool POST lên cloud. Server sẽ archive
         version cũ, save new, rebuild manifest, append audit log.

    Sau khi publish OK, các máy khác chạy vbhc_sync_knowledge sẽ nhận version mới.

    Args:
        slug: slug loại VB (vd "bao-cao", "cong-van"). Phải có file
              ~/.vbhc/cache/templates/<slug>.docx (do vbhc_update_template ghi).
        confirmed: True khi user xác nhận publish (lần gọi thứ 2).

    Returns:
        dict gồm: slug, source, size_local, sha256_local, confirmed, applied,
                  preview_message hoặc result (response từ server: sha256,
                  archived_to, manifest_generated).
    """
    if not kc.is_configured():
        return {"error": "Chưa bootstrap — chạy `python mcp/bootstrap.py` trước."}

    safe_slug = re.sub(r"[^a-z0-9\-_]", "", slug.lower())
    if not safe_slug:
        return {"error": f"slug không hợp lệ: '{slug}'"}

    cache_p = kc.cache_path_for("templates", f"{safe_slug}.docx")
    if not cache_p.is_file():
        return {
            "error": (
                f"Không có {cache_p}. Chạy vbhc_update_template để ghi template "
                f"vào cache trước, rồi publish."
            )
        }

    body = cache_p.read_bytes()
    sha256_local = __import__("hashlib").sha256(body).hexdigest()

    result: dict[str, Any] = {
        "slug": safe_slug,
        "source": str(cache_p),
        "size_local": len(body),
        "sha256_local": sha256_local,
        "confirmed": confirmed,
        "applied": False,
    }

    if not confirmed:
        result["preview_message"] = (
            f"Sẽ POST '{cache_p.name}' ({len(body)} bytes, sha256={sha256_local[:12]}...) "
            f"lên cloud. Server sẽ archive version cũ + rebuild manifest. "
            f"Gọi lại với confirmed=True để thực hiện."
        )
        return result

    try:
        resp = kc.publish_template(safe_slug, source_path=cache_p)
    except kc.KBError as e:
        result["error"] = str(e)
        return result

    result["applied"] = True
    result["result"] = resp
    # Sau publish, refresh local manifest để vbhc_knowledge_status thấy đúng
    try:
        kc.sync_manifest()
    except kc.KBError:
        pass
    result["message"] = (
        f"✓ Đã publish '{safe_slug}.docx' lên cloud "
        f"(sha256={resp.get('sha256','?')[:12]}..., "
        f"archived={'yes' if resp.get('archived_to') else 'no'}). "
        f"Máy khác chạy vbhc_sync_knowledge sẽ nhận version mới."
    )
    return result


@mcp.tool()
def vbhc_knowledge_status() -> dict:
    """Trả thông tin cache: bản local vs cloud, drift, last sync.

    Returns:
        dict với: configured, cloud_url, cached_assets (templates/rules có
        trong cache), local_manifest (versions), cloud_manifest (nếu kết
        nối được), drift (asset cần sync), last_sync (summary lần sync trước).
    """
    return kc.status()


# =====================================================================
# Entry point: stdio transport (local thin-MCP)
# =====================================================================
# HTTP mode đã bị gỡ ở v1.0. v0.9 cũ có `--http` chạy như cloud MCP với
# auth Bearer per-client; v1.0 đổi sang kiến trúc local thin-MCP + cloud
# Knowledge Hub (cloud/kb_server.py). Migration: xem MIGRATION-v1.0.md.

if __name__ == "__main__":
    mcp.run()
